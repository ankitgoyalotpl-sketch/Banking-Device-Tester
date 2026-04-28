from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def create_manual():
    doc = Document()

    # Document Title & Subtitle
    title = doc.add_heading('Oxymora GPS Device Tester v8.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Comprehensive User & Hardware Integration Manual')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    doc.add_paragraph('The Oxymora GPS Device Tester is a secure, web-based diagnostic platform designed explicitly for '
                      'authenticating, testing, and troubleshooting GPS tracking hardware provisioned for financial institutions. '
                      'Built on a modern React architecture, this software interfaces directly with hardware modules over USB '
                      'using advanced Web Serial APIs, bypassing the need for native desktop applications.')

    add_heading(doc, '1.1 Supported Hardware Modules', 2)
    doc.add_paragraph('• State Bank of India (SBI) - Proprietary Cipher Authentication Hardware', style='List Bullet')
    doc.add_paragraph('• Bank of Baroda (BOB) - Advanced JSON Telemetry Hardware', style='List Bullet')
    doc.add_paragraph('• Bank of Maharashtra (BOM) - Legacy Command Hardware', style='List Bullet')

    # 2. System Prerequisites
    add_heading(doc, '2. System Prerequisites', 1)
    doc.add_paragraph('Ensure your workstation meets the following requirements before operating the portal:')
    doc.add_paragraph('1. Web Browser: Google Chrome (v89+), Microsoft Edge (v89+), or Opera. Safari and Firefox are currently unsupported due to Web Serial API limitations.')
    doc.add_paragraph('2. USB Drivers: Ensure CH340, CP2102, or PL2303 UART-to-USB drivers are installed depending on your specific GPS dongle.')
    doc.add_paragraph('3. Network: Active broadband connection to sync payloads with the central Vercel/Firebase backend.')

    doc.add_page_break()

    # 3. User Interface & Initial Setup
    add_heading(doc, '3. Getting Started: The Interface', 1)
    add_heading(doc, '3.1 Dashboard Landing & Bank Selection', 2)
    doc.add_paragraph('Upon navigating to the portal URL, you are greeted by the primary Bank Selection interface. '
                      'The system isolates testing protocols based on the target bank to prevent cryptographic mismatches.')
    
    img1 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777357109118.png"
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 1: Main Landing Page', style='Caption')

    add_heading(doc, '3.2 Security & Compliance (T&C)', 2)
    doc.add_paragraph('Due to the sensitive nature of tracking telemetry, operators must consent to the Data Collection and Privacy Policy. Check the consent box to proceed.')
    
    img2 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777359852316.png"
    if os.path.exists(img2):
        doc.add_picture(img2, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 4. Hardware Connectivity
    add_heading(doc, '4. Hardware Connectivity via Web Serial', 1)
    doc.add_paragraph('1. Connect your GPS hardware dongle to an available USB port.')
    doc.add_paragraph('2. Navigate to your target bank\'s testing suite and locate the blue "CONNECT DEVICE" button.')
    
    img3 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777359900948.png"
    if os.path.exists(img3):
        doc.add_picture(img3, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('3. A native browser security prompt will appear. Select your paired COM port (e.g., COM16) and click "Connect".')
    
    img4 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777359932166.png"
    if os.path.exists(img4):
        doc.add_picture(img4, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 5. Diagnostic Testing
    add_heading(doc, '5. Running Diagnostic Tests', 1)
    doc.add_paragraph('After successful connection, the indicator badge will switch to "Connected". Click "Run Device Test" to dispatch the command payload.')
    
    add_heading(doc, '5.1 BOB Device Testing', 2)
    doc.add_paragraph('The application will dispatch a JSON structure requesting hardware specifications. It will render the Firmware Version, Serial Number, and automatically trigger a background location poll to extract Live GPS Coordinates.')

    add_heading(doc, '5.2 SBI Device Testing', 2)
    doc.add_paragraph('SBI modules utilize proprietary cryptography. The portal will dispatch the SBI_MAGIC_STRING. If the hardware authenticates it, a "Success" flag is raised, verifying the hardware\'s integrity.')
    
    img5 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777360007588.png"
    if os.path.exists(img5):
        doc.add_picture(img5, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 6. Additional Tools
    add_heading(doc, '6. Support, Ticketing, and Settings', 1)
    add_heading(doc, '6.1 Helpdesk & Complaints', 2)
    doc.add_paragraph('If a hardware dongle fails authentication, you can raise an internal ticket directly via the "Create Complaint" tab. Fill in your Device Serial Number and a description of the error trace.')
    
    img7 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777360161691.png"
    if os.path.exists(img7):
        doc.add_picture(img7, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Alternatively, use the "Contact" button on the navigation bar to fetch live support details.')
    
    img8 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777360187307.png"
    if os.path.exists(img8):
        doc.add_picture(img8, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, '6.2 UI Accessibility', 2)
    doc.add_paragraph('Users can toggle between Light and Dark mode using the sun/moon icon at the top right, accommodating varying workplace lighting conditions.')

    img6 = r"C:\Users\ankit\.gemini\antigravity\brain\e355ddda-385c-4ea5-9c1e-9b3bf3f331c7\media__1777360034891.png"
    if os.path.exists(img6):
        doc.add_picture(img6, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Troubleshooting
    add_heading(doc, '7. Troubleshooting', 1)
    doc.add_paragraph('• No COM Port Found: Ensure the USB cable supports data transfer (not charging-only). Verify CH340 drivers are installed via Windows Device Manager.')
    doc.add_paragraph('• Invalid Response Error: Disconnect and reconnect the device. If the issue persists, the hardware may be bound to a different banking protocol.')

    save_path = r"C:\Users\ankit\OneDrive\Desktop\Oxymora_GPS_Tester_Detailed_Manual.docx"
    doc.save(save_path)
    print(f"Document saved successfully to {save_path}")

if __name__ == "__main__":
    create_manual()
